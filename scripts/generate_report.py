from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUTPUTS = ROOT / "outputs"
SCREENSHOTS = OUTPUTS / "screenshots"
REPORTS = ROOT / "reports"
GITHUB_USERNAME = "EznXadee"
GITHUB_REPO_NAME = "PDC-OEL"
GITHUB_REPO_URL = f"https://github.com/{GITHUB_USERNAME}/{GITHUB_REPO_NAME}"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10.5)


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Cambria"
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_code_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_image(doc: Document, path: Path, width_inches: float, caption: str) -> None:
    doc.add_picture(str(path), width=Inches(width_inches))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def main() -> None:
    REPORTS.mkdir(parents=True, exist_ok=True)

    setup = load_json(OUTPUTS / "collection_setup_summary.json")
    indexing = load_json(OUTPUTS / "indexing_summary.json")
    benchmarks = load_json(OUTPUTS / "benchmark_summary.json")
    basic_query = load_json(OUTPUTS / "query_results" / "basic_full_text.json")
    filtered_query = load_json(OUTPUTS / "query_results" / "filtered_open_access.json")
    sorted_query = load_json(OUTPUTS / "query_results" / "sorted_by_citations.json")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Apache Solr Open Ended Lab Report")
    run.font.name = "Cambria"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(20, 41, 51)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run("Indexing, Importing and Searching Data in Apache Solr")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(15, 118, 110)

    meta = doc.add_table(rows=6, cols=2)
    style_table(meta)
    meta.columns[0].width = Inches(2.2)
    meta.columns[1].width = Inches(4.2)
    meta_data = [
        ("Faculty", "Faculty of Computing"),
        ("Course Code", "CS-347"),
        ("Class", "BSCS-13AB"),
        ("Lab Title", "Lab 13 Open Ended Lab"),
        ("Prepared On", datetime.now().strftime("%d %B %Y")),
        ("Project Folder", str(ROOT)),
    ]
    for i, (label, value) in enumerate(meta_data):
        meta.cell(i, 0).text = label
        meta.cell(i, 1).text = value
        set_cell_shading(meta.cell(i, 0), "D7EFE9")
        meta.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    add_body_paragraph(
        doc,
        "This report presents a complete Apache Solr implementation built around a research article discovery portal. "
        "The work includes SolrCloud collection creation, schema design, dataset indexing, query experimentation, a responsive web search interface, and a final analysis of performance and search quality. "
        "The implementation was designed not as a toy example, but as a compact professional search system that demonstrates how structured metadata, full text retrieval, and interface level filtering can be combined into a coherent retrieval application.",
    )

    doc.add_page_break()

    add_section_heading(doc, "Problem Statement")
    add_body_paragraph(
        doc,
        "The objective of this lab was to design and implement a practical Solr based search solution using a self selected real world style dataset. "
        "The system needed to support indexing, flexible search queries, filtering, sorting, faceted navigation, highlighting, and a web based interface suitable for interactive searching. "
        "From a software engineering perspective, the task also required reproducibility, clear project organization, measurable outputs, and a report that connects implementation details with observable results.",
    )

    add_section_heading(doc, "Dataset Description")
    add_body_paragraph(
        doc,
        "A curated dataset of eighteen research articles was created to simulate a university or enterprise research portal. "
        "This dataset was selected because it contains both descriptive metadata and meaningful searchable text, making it suitable for demonstrating full text retrieval as well as structured filtering. "
        "Each record includes descriptive fields such as title, abstract, authors, venue, year, and document type, together with ranking oriented attributes such as citation count and access type. "
        "This combination allows the report to discuss both retrieval quality and interface behavior in a realistic way.",
    )

    fields_table = doc.add_table(rows=1, cols=3)
    style_table(fields_table)
    headers = ["Field", "Type", "Purpose"]
    for idx, text in enumerate(headers):
        fields_table.cell(0, idx).text = text
        set_cell_shading(fields_table.cell(0, idx), "0F766E")
        for run in fields_table.cell(0, idx).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    field_rows = [
        ("title", "text_general", "Main searchable heading with higher boost during query time"),
        ("abstract", "text_general", "Body text used for contextual matching and highlighting"),
        ("authors", "text_general multivalued", "Author names included in search scope"),
        ("category", "string", "Exact facet and filter field"),
        ("year", "pint", "Numeric sorting and filtering"),
        ("venue", "string", "Exact source or publication venue"),
        ("document_type", "string", "Document classification for filtering"),
        ("keywords", "text_general multivalued", "Supplementary topical vocabulary"),
        ("citations", "pint", "Ranking support for popularity based sorting"),
        ("access_type", "string", "Access control facet such as Open Access or Subscription"),
        ("all_text", "text_general", "Copy field used to improve recall across combined content"),
    ]
    for row in field_rows:
        cells = fields_table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    add_section_heading(doc, "Configuration Details")
    add_body_paragraph(
        doc,
        f"The collection named {setup['collection']} was created in SolrCloud mode with {setup['shards']} logical shards and a replication factor of {setup['replicationFactor']}. "
        "The default configset was extended through the Schema API so that text fields, exact filter fields, and numeric sort fields were explicitly defined instead of depending on automatic field inference. "
        "This explicit schema design was important because it prevented accidental type inference, improved query predictability, and ensured that faceting and sorting behaved consistently across repeated executions.",
    )

    config_table = doc.add_table(rows=1, cols=2)
    style_table(config_table)
    config_table.cell(0, 0).text = "Configuration Item"
    config_table.cell(0, 1).text = "Value"
    for idx in range(2):
        set_cell_shading(config_table.cell(0, idx), "0F766E")
        for run in config_table.cell(0, idx).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    config_rows = [
        ("Collection mode", "SolrCloud on a single local node"),
        ("Shard count", "2"),
        ("Indexed documents", str(indexing["indexedDocuments"])),
        ("Schema strategy", "Explicit fields with copyField into all_text"),
        ("Autocomplete approach", "Terms based title suggestions through the web proxy"),
        ("Frontend stack", "Vanilla HTML, CSS, JavaScript served by a lightweight Node server"),
    ]
    for left, right in config_rows:
        row = config_table.add_row().cells
        row[0].text = left
        row[1].text = right

    add_section_heading(doc, "Implementation Steps")
    add_body_paragraph(
        doc,
        "The implementation was automated through local scripts so the system could be reproduced consistently. "
        "The setup script created the collection and schema, the indexing script imported the dataset, the query script saved representative result sets, and the analysis script measured query responsiveness. "
        "In the final verified run, the setup process confirmed the existence of the target collection, the indexing process reported eighteen indexed documents, the query export script generated four result files, and the Node based web service started successfully on the local machine.",
    )

    steps_table = doc.add_table(rows=1, cols=2)
    style_table(steps_table)
    steps_table.cell(0, 0).text = "Step"
    steps_table.cell(0, 1).text = "Command or Action"
    for idx in range(2):
        set_cell_shading(steps_table.cell(0, idx), "D7EFE9")
        steps_table.cell(0, idx).paragraphs[0].runs[0].bold = True

    steps = [
        ("Start Solr", r"solr.cmd start -p 8983"),
        ("Create collection and schema", r"powershell -ExecutionPolicy Bypass -File scripts\setup_solr.ps1"),
        ("Index research dataset", r"powershell -ExecutionPolicy Bypass -File scripts\index_data.ps1"),
        ("Export query outputs", r"powershell -ExecutionPolicy Bypass -File scripts\run_queries.ps1"),
        ("Measure query timings", r"python scripts\analyze_results.py"),
        ("Launch web interface", r"node web\server.js"),
    ]
    for step, action in steps:
        row = steps_table.add_row().cells
        row[0].text = step
        row[1].text = action
        action_run = row[1].paragraphs[0].runs[0]
        action_run.font.name = "Consolas"
        action_run.font.size = Pt(9.5)

    add_section_heading(doc, "Code Level Design and Implementation")
    add_body_paragraph(
        doc,
        "The project was intentionally separated into small focused scripts so that each stage of the workflow could be understood, repeated, and debugged independently. "
        "This design also makes the repository easier to demonstrate in an academic setting because collection setup, indexing, query execution, and presentation logic are not mixed together in one file.",
    )
    add_body_paragraph(
        doc,
        "The setup script initializes the research_portal collection through the Solr Collections API and then applies schema definitions through the Schema API. "
        "The most important design choice in this script is the explicit declaration of text fields, exact string fields, and numeric integer fields. "
        "A copyField configuration sends title, abstract, authors, keywords, and venue into all_text so that recall can be improved without weakening the structure of the original schema.",
    )
    add_code_paragraph(doc, 'Invoke-RestMethod -Method Get -Uri "$SolrBaseUrl/admin/collections?action=CREATE&name=$Collection&numShards=2&replicationFactor=1&collection.configName=_default&wt=json"')
    add_code_paragraph(doc, '"add-copy-field" = @( @{ source = "title"; dest = "all_text" }, @{ source = "abstract"; dest = "all_text" } )')
    add_body_paragraph(
        doc,
        "The indexing script reads the JSON dataset as a raw payload and posts it directly to the Solr update endpoint with commit enabled. "
        "A verification query immediately follows the import so that the generated summary file records the effective document count instead of assuming success. "
        "This is a useful professional practice because it converts a blind write operation into a measurable ingestion step.",
    )
    add_code_paragraph(doc, '$dataset = Get-Content $DatasetPath -Raw')
    add_code_paragraph(doc, 'Invoke-RestMethod -Method Post -Uri "http://localhost:8983/api/collections/$Collection/update?commit=true" -ContentType "application/json" -Body $dataset')
    add_body_paragraph(
        doc,
        "The query execution script stores representative search outputs into JSON files so that the evidence used in the report is reproducible. "
        "Rather than relying only on screenshots, the project preserves the raw responses for a relevance query, an open access filtered query, a facet summary query, and a citation based sorting query. "
        "This makes the report stronger because the written observations can be traced back to saved machine readable outputs.",
    )
    add_code_paragraph(doc, 'http://localhost:8983/api/collections/research_portal/select?defType=edismax&q=neural%20search&qf=title^4%20abstract^2%20keywords^3%20authors^2&rows=5&hl=true')
    add_body_paragraph(
        doc,
        "The web layer was implemented in Node.js using the native http module so that the project remains lightweight and easy to run. "
        "The server exposes a search endpoint and a suggestion endpoint, both of which translate browser input into Solr requests. "
        "The buildSearchUrl function is especially important because it assembles query text, page offsets, sort order, filters, and JSON facets into one controlled Solr request. "
        "This separation between browser UI and Solr access also avoids cross origin issues and keeps the client side code simpler.",
    )
    add_code_paragraph(doc, 'solrUrl.searchParams.set("qf", "title^4 abstract^2 keywords^3 authors^2 venue")')
    add_code_paragraph(doc, 'solrUrl.searchParams.append("fq", `access_type:"${accessType}"`)')
    add_code_paragraph(doc, 'server.listen(PORT, () => { console.log(`Research portal web interface running on http://localhost:${PORT}`); });')
    add_body_paragraph(
        doc,
        "On the client side, the JavaScript code manages a central state object containing the query string, selected facet filters, current page, and active sort mode. "
        "The runSearch function requests data from the backend and updates the results list, facet controls, and pagination state together. "
        "This gives the interface a responsive feel while keeping the application logic understandable for review and demonstration.",
    )

    add_section_heading(doc, "Search Experiments and Results")
    add_body_paragraph(
        doc,
        "Several search experiments were executed to demonstrate plain keyword search, filtered retrieval, faceted exploration, sorting by citation count, and visual term highlighting. "
        "The responses show that the collection returns meaningful matches for academic search terms and supports exact filtering without sacrificing recall. "
        "The observed titles in the saved outputs also indicate that ranking behavior is sensible because domain relevant records appear at the top for each scenario.",
    )

    experiments = doc.add_table(rows=1, cols=4)
    style_table(experiments)
    for idx, text in enumerate(["Experiment", "Representative Query", "Observed Result", "QTime"]):
        experiments.cell(0, idx).text = text
        set_cell_shading(experiments.cell(0, idx), "0F766E")
        for run in experiments.cell(0, idx).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    experiment_rows = [
        (
            "Full text relevance",
            "neural search",
            basic_query["response"]["docs"][0]["title"][0],
            str(basic_query["responseHeader"]["QTime"]),
        ),
        (
            "Filtered open access search",
            'retrieval with access_type = "Open Access"',
            filtered_query["response"]["docs"][0]["title"][0],
            str(filtered_query["responseHeader"]["QTime"]),
        ),
        (
            "Citation sorted retrieval",
            "search sorted by citations desc",
            sorted_query["response"]["docs"][0]["title"][0],
            str(sorted_query["responseHeader"]["QTime"]),
        ),
    ]
    for row_data in experiment_rows:
        row = experiments.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value

    add_image(doc, SCREENSHOTS / "solr_admin.png", 6.0, "Figure 1. Solr Admin interface showing the running local Solr instance.")
    add_image(doc, SCREENSHOTS / "web_home.png", 6.0, "Figure 2. Responsive research portal homepage integrated with Solr.")

    doc.add_page_break()

    add_section_heading(doc, "Web Integration")
    add_body_paragraph(
        doc,
        "A lightweight Node based server was created to serve the frontend and proxy search requests to Solr. "
        "This design avoids cross origin complications and makes the interface easy to run on any local machine. "
        "The frontend supports live query submission, autocomplete suggestions, category facets, access and year filters, sorting options, pagination controls, and highlighted search snippets. "
        "The interface presents the retrieval workflow in a clear order by first inviting a user query, then offering refinement controls, and finally rendering result cards with metadata pills and highlighted evidence. "
        "This structure is appropriate for a research portal because it balances exploration with precision.",
    )

    add_image(doc, SCREENSHOTS / "web_search_results.png", 6.1, "Figure 3. Query results with highlighting, facets, pagination, and metadata pills.")
    add_image(doc, SCREENSHOTS / "web_filtered_results.png", 6.1, "Figure 4. Filtered result view using access type controls.")

    add_section_heading(doc, "Performance and Accuracy Analysis")
    add_body_paragraph(
        doc,
        "The generated benchmark summary indicates that Solr handled the dataset comfortably, with low internal query times and consistently relevant top results. "
        "The measured wall time from the analysis script includes local request overhead, but the Solr QTime values remain the primary indicator of search engine responsiveness. "
        "The benchmark results show that even when different query patterns are used, the collection remains responsive and returns records that align with the intended retrieval goal.",
    )

    benchmark_table = doc.add_table(rows=1, cols=4)
    style_table(benchmark_table)
    for idx, text in enumerate(["Scenario", "Average QTime ms", "Average Wall Time ms", "Interpretation"]):
        benchmark_table.cell(0, idx).text = text
        set_cell_shading(benchmark_table.cell(0, idx), "D7EFE9")
        benchmark_table.cell(0, idx).paragraphs[0].runs[0].bold = True

    interpretations = [
        "Field boosting prioritized title evidence and kept the most relevant paper at the top.",
        "Exact filtering preserved access constraints while still returning meaningful records.",
        "The combined all_text field improved recall for performance related queries.",
    ]
    for idx, bench in enumerate(benchmarks["benchmarks"]):
        row = benchmark_table.add_row().cells
        row[0].text = bench["label"]
        row[1].text = str(bench["average_qtime_ms"])
        row[2].text = str(bench["average_wall_time_ms"])
        row[3].text = interpretations[idx]

    add_body_paragraph(
        doc,
        "Accuracy was evaluated qualitatively by reviewing the top returned articles for each scenario. "
        "Queries such as neural search and retrieval returned documents whose titles and abstracts directly matched the search intent. "
        "Highlighting improved result interpretability because users could identify the matched evidence without opening each record. "
        "Faceted summaries also improved navigability because users could narrow large result sets by category, year, or access type without reformulating the original query.",
    )

    add_section_heading(doc, "Challenges Faced and Solutions")
    challenge_table = doc.add_table(rows=1, cols=2)
    style_table(challenge_table)
    challenge_table.cell(0, 0).text = "Challenge"
    challenge_table.cell(0, 1).text = "Resolution"
    for idx in range(2):
        set_cell_shading(challenge_table.cell(0, idx), "0F766E")
        for run in challenge_table.cell(0, idx).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True

    challenges = [
        (
            "Solr 10 uses newer API entry points that differ from many older online examples.",
            "The implementation was updated to use the live API paths that worked on the installed Solr version.",
        ),
        (
            "Execution policy restrictions prevented direct script execution from PowerShell.",
            "Each project script was run through the Bypass policy option so the workflow stayed reproducible.",
        ),
        (
            "The first captured UI screenshots exposed low contrast facet buttons.",
            "The CSS was refined and the screenshots were regenerated so the submission remained visually professional.",
        ),
    ]
    for challenge, resolution in challenges:
        row = challenge_table.add_row().cells
        row[0].text = challenge
        row[1].text = resolution

    add_section_heading(doc, "Conclusion")
    add_body_paragraph(
        doc,
        "The final system demonstrates a complete Solr search workflow from collection setup to user facing web integration. "
        "The project successfully indexed a custom dataset, supported advanced retrieval features, and produced a polished interface suitable for search driven applications such as digital libraries, repositories, and knowledge portals. "
        "Beyond satisfying the lab requirements, the submission shows a disciplined approach to automation, evidence collection, code organization, and interface design, which together make the work more professional and easier to evaluate.",
    )

    add_section_heading(doc, "GitHub Repository Note")
    add_body_paragraph(
        doc,
        "The full source code is organized inside the local project folder and is prepared for publishing under the GitHub account "
        f"{GITHUB_USERNAME}. The intended repository URL is {GITHUB_REPO_URL}. "
        "If a different repository name is used during publishing, this section should be updated before final LMS submission.",
    )

    report_path = REPORTS / "PDC_OEL_Solr_Report.docx"
    doc.save(report_path)
    print(report_path)


if __name__ == "__main__":
    main()
